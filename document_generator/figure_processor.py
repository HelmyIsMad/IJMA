"""
Figure and image processing utilities for document generation.
Handles base64 image decoding, image insertion, and figure formatting.
"""
from typing import Optional
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
import re
import base64
import io
from .style_handler import apply_hardcoded_style


def process_image_content(paragraph: Paragraph, html_content: str, figure_num: int = 1, caption: str = "") -> None:
    """
    Process HTML content containing images and add them to the paragraph.
    
    Args:
        paragraph: The paragraph to add images to
        html_content: HTML content with base64-encoded images
        figure_num: Figure number for caption
        caption: Figure caption text
    """
    try:
        # Extract image data from HTML
        img_pattern = r'<img[^>]+src=["\'](data:image/[^;]+;base64,[^"\']+)["\'][^>]*>'
        matches = re.findall(img_pattern, html_content, re.IGNORECASE)
        
        if matches:
            for img_src in matches:
                try:
                    _process_base64_image(paragraph, img_src)
                except Exception as e:
                    # If image processing fails, add as text
                    run = paragraph.add_run(f"[Image could not be processed: {str(e)}]")
                    apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
        else:
            # No images found, treat as regular text but strip HTML tags
            _fallback_to_text(paragraph, html_content)
                
    except Exception as e:
        # Fallback: add as text
        run = paragraph.add_run(f"[Figure processing error: {str(e)}]")
        apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
        paragraph.add_run("\n")


def _process_base64_image(paragraph: Paragraph, img_src: str) -> None:
    """
    Decode and insert a base64-encoded image into the paragraph.
    
    Args:
        paragraph: The paragraph to add the image to
        img_src: Base64-encoded image data URL
    """
    # Parse the data URL
    if img_src.startswith('data:image/'):
        # Extract format and base64 data
        header, base64_data = img_src.split(',', 1)
        img_format = header.split('/')[1].split(';')[0]
        
        # Decode base64 image data
        image_data = base64.b64decode(base64_data)
        image_stream = io.BytesIO(image_data)
        
        # Add image to the paragraph
        run = paragraph.add_run()
        pic = run.add_picture(image_stream, width=Inches(6.5))  # Page width
        
        # Set text wrapping and add border
        _apply_image_wrapping_and_border(pic)


def _apply_image_wrapping_and_border(pic) -> None:
    """
    Apply text wrapping (top and bottom) and 1px border to an image.
    
    Args:
        pic: The picture object to apply wrapping to
    """
    try:
        # Get the inline element and convert to anchor for wrapping
        inline = pic._inline
        
        # Create anchor element for wrapping (instead of inline)
        anchor = OxmlElement('wp:anchor')
        
        # Copy necessary attributes from inline
        anchor.set('distT', '0')
        anchor.set('distB', '0')
        anchor.set('distL', '114300')
        anchor.set('distR', '114300')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '251658240')
        anchor.set('behindDoc', '0')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        
        # Simple positioning
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        anchor.append(simplePos)
        
        # Position horizontal (center to page)
        positionH = OxmlElement('wp:positionH')
        positionH.set('relativeFrom', 'column')
        align = OxmlElement('wp:align')
        align.text = 'center'
        positionH.append(align)
        anchor.append(positionH)
        
        # Position vertical (relative to paragraph)
        positionV = OxmlElement('wp:positionV')
        positionV.set('relativeFrom', 'paragraph')
        posOffset = OxmlElement('wp:posOffset')
        posOffset.text = '0'
        positionV.append(posOffset)
        anchor.append(positionV)
        
        # Copy extent (size)
        extent = inline.find(qn('wp:extent'))
        if extent is not None:
            anchor.append(extent)
        
        # Effect extent
        effectExtent = OxmlElement('wp:effectExtent')
        effectExtent.set('l', '0')
        effectExtent.set('t', '0')
        effectExtent.set('r', '0')
        effectExtent.set('b', '0')
        anchor.append(effectExtent)
        
        # Wrapping - topAndBottom
        wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
        anchor.append(wrapTopAndBottom)
        
        # Document properties
        docPr = inline.find(qn('wp:docPr'))
        if docPr is not None:
            anchor.append(docPr)
        
        # Copy graphic and add border
        graphic = inline.find(qn('a:graphic'))
        if graphic is not None:
            anchor.append(graphic)
            _add_image_border(graphic)
        
        # Replace inline with anchor
        inline.getparent().replace(inline, anchor)
        
    except Exception as e:
        # If wrapping/border fails, continue without it
        pass


def _add_image_border(graphic) -> None:
    """
    Add a 1px black border to an image.
    
    Args:
        graphic: The graphic element to add border to
    """
    try:
        graphicData = graphic.graphicData
        pic_element = graphicData.pic
        spPr = pic_element.spPr
        
        # Create border outline
        if spPr.find(qn('a:ln')) is None:
            ln = OxmlElement('a:ln')
            ln.set('w', '9525')  # 1px in EMUs (English Metric Units)
            
            # Create solid fill for border
            solidFill = OxmlElement('a:solidFill')
            srgbClr = OxmlElement('a:srgbClr')
            srgbClr.set('val', '000000')  # Black border
            solidFill.append(srgbClr)
            ln.append(solidFill)
            
            spPr.append(ln)
    except:
        pass


def _fallback_to_text(paragraph: Paragraph, html_content: str) -> None:
    """
    Fallback: strip HTML tags and add as regular text.
    
    Args:
        paragraph: The paragraph to add text to
        html_content: HTML content to convert to text
    """
    import html
    clean_text = re.sub(r'<[^>]+>', '', html_content)
    clean_text = html.unescape(clean_text)
    
    if clean_text.strip():
        run = paragraph.add_run(clean_text)
        apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
