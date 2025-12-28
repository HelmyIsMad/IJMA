"""
Value insertion utilities for document placeholders.
Handles replacement of template variables with formatted content.
"""
from typing import List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from .config import VALUES
from .style_handler import apply_hardcoded_style
from .text_formatter import process_formatted_text
from .table_processor import process_table_content
from .figure_processor import process_image_content

def int_percent_to_float_percent(s):
    return re.sub(
        r'(?<!\.)\d+(?=%)',
        lambda m: f"{m.group(0)}.0",
        s
    )

def normalize_symbol_spacing(text: str) -> str:
    """
    Ensure "=" and "±" have exactly one space before and after them.
    
    Args:
        text: The text to normalize
        
    Returns:
        Text with normalized spacing around = and ± symbols
    """
    # Remove all spaces around = and ±
    text = re.sub(r'\s*=\s*', '=', text)
    text = re.sub(r'\s*±\s*', '±', text)
    
    # Add exactly one space before and after
    text = re.sub(r'=', ' = ', text)
    text = re.sub(r'±', ' ± ', text)
    
    # Clean up any double spaces that might have been created
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def insert_values(paragraph: Paragraph, variable: str, doc: Optional[Document] = None) -> None:
    """
    Insert formatted values into a paragraph based on the variable placeholder.
    
    Args:
        paragraph: The paragraph to insert values into
        variable: The placeholder variable name (e.g., "{{title}}")
        doc: The document object (needed for tables/figures)
        paragraphs: List of all paragraphs (unused, kept for compatibility)
    """
    # Clear the paragraph content
    paragraph.clear()
    
    if variable == "{{header_name}}":
        _insert_header_name(paragraph)
    elif variable == "{{research_type}}":
        _insert_research_type(paragraph)
    elif variable == "{{research_title}}":
        _insert_research_title(paragraph)
    elif variable == "{{authors}}":
        _insert_authors(paragraph)
    elif variable == "{{affiliation}}":
        _insert_affiliation(paragraph)
    elif variable == "{{date_received}}" or variable == "{{date_accepted}}":
        _insert_date(paragraph, variable)
    elif variable == "{{email}}":
        _insert_email(paragraph)
    elif variable == "{{citation}}":
        _insert_citation(paragraph)
    elif variable == "{{abstract}}":
        _insert_abstract(paragraph)
    elif variable == "{{keywords}}":
        _insert_keywords(paragraph)
    elif variable in ["{{intro}}", "{{aim}}", "{{methods}}", "{{results}}", "{{discussion}}", "{{references}}"]:
        _insert_content_section(paragraph, variable)
    elif variable == "{{tables}}":
        _insert_tables(paragraph, doc)
    elif variable == "{{figures}}":
        _insert_figures(paragraph, doc)


def _insert_header_name(paragraph: Paragraph) -> None:
    """Insert header name with special formatting."""
    color = "#1F3864"
    
    run1 = paragraph.add_run(VALUES["{{header_name}}"][0])
    apply_hardcoded_style(run1, font_name="Arial Narrow", font_size=11, bold=True, color=color, spacing=0, scale=100)
    
    run2 = paragraph.add_run(VALUES["{{header_name}}"][1])
    apply_hardcoded_style(run2, font_name="Arial Narrow", font_size=11, bold=True, italic=True, color=color, spacing=0, scale=100)


def _insert_research_type(paragraph: Paragraph) -> None:
    """Insert research type in red."""
    run1 = paragraph.add_run(VALUES["{{research_type}}"][0])
    apply_hardcoded_style(run1, font_name="Times New Roman", font_size=13, color="#FF0000", spacing=0, scale=100)


def _insert_research_title(paragraph: Paragraph) -> None:
    """Insert research title with percentage formatting."""
    title_text = int_percent_to_float_percent(VALUES["{{research_title}}"][0])
    title_text = normalize_symbol_spacing(title_text)
    run1 = paragraph.add_run(title_text)
    apply_hardcoded_style(run1, font_name="Times New Roman", font_size=15, bold=True, spacing=0, scale=100)


def _insert_authors(paragraph: Paragraph) -> None:
    """Insert authors with superscript markers.

    Supports tokens in VALUES['{{authors}}'] like:
      - ["Alice", "*", ";", "Bob", ";"]
      - ["Alice", "*1", ";", "Bob", "2", ";"]

    i.e., author name first, then optional marker token, then semicolon.
    """
    color2 = "#FF0000"

    def _is_marker_token(t: str) -> bool:
        t = (t or '').strip()
        if not t:
            return False
        if t == '*':
            return True
        if t.isdigit():
            return True
        # '*1', '*2', etc.
        if t.startswith('*') and t[1:].isdigit():
            return True
        return False

    for token in VALUES["{{authors}}"]:
        formatted = int_percent_to_float_percent(str(token))
        run = paragraph.add_run(formatted)

        if str(token) == ';':
            apply_hardcoded_style(run, font_name="Times New Roman", font_size=10, bold=True, spacing=0, scale=100)
        elif _is_marker_token(str(token)):
            apply_hardcoded_style(
                run,
                font_name="Times New Roman",
                font_size=10,
                bold=True,
                color=color2,
                superscript=True,
                spacing=0,
                scale=100,
            )
        else:
            # Author name
            apply_hardcoded_style(run, font_name="Times New Roman", font_size=10, bold=True, spacing=0, scale=100)


def _insert_affiliation(paragraph: Paragraph) -> None:
    """Insert affiliations.

    Supports two formats:
    - Numbered list tokens: ["1", "Aff one.", "2", "Aff two."]
    - Single/un-numbered: ["Aff one."]
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    items = VALUES["{{affiliation}}"]

    # Un-numbered mode if the first token isn't a digit
    if not items:
        return

    if not str(items[0]).strip().isdigit():
        # Print each item on its own line
        for idx, aff in enumerate(items):
            run = paragraph.add_run(str(aff))
            apply_hardcoded_style(run, font_size=8, spacing=0, scale=100)
            if idx < len(items) - 1:
                paragraph.add_run("\n")
        return

    # Numbered mode: tokens are [num, aff, num, aff, ...]
    for i, text in enumerate(items):
        run = paragraph.add_run(str(text))
        if i % 2:
            apply_hardcoded_style(run, font_size=8, spacing=0, scale=100)
            paragraph.add_run("\n")
        else:
            apply_hardcoded_style(run, font_size=8, superscript=True, color="#FF0000", spacing=0, scale=100)
            paragraph.add_run(" ")


def _insert_date(paragraph: Paragraph, variable: str) -> None:
    """Insert date."""
    run = paragraph.add_run(VALUES[variable][0])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=10, spacing=0, scale=100)


def _insert_email(paragraph: Paragraph) -> None:
    """Insert email with label."""
    run = paragraph.add_run(VALUES["{{email}}"][0])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=10, bold=True, spacing=0, scale=100)
    
    run = paragraph.add_run(VALUES["{{email}}"][1])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=10, spacing=0, scale=100)


def _insert_citation(paragraph: Paragraph) -> None:
    """Insert citation."""
    run = paragraph.add_run(VALUES["{{citation}}"][0])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=8.5, color="#FF0000", bold=True, spacing=0, scale=100)
    
    run = paragraph.add_run(VALUES["{{citation}}"][1])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=9, spacing=0, scale=100)


def _insert_abstract(paragraph: Paragraph) -> None:
    """Insert abstract sections with bold headers."""
    current_p_element = paragraph._p
    current_paragraph = paragraph
    
    # Process abstract sections - each section (header + content) in one paragraph
    for i in range(0, len(VALUES["{{abstract}}"]), 2):
        if i + 1 < len(VALUES["{{abstract}}"]):
            # Set justified alignment for the paragraph
            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Set paragraph spacing - 6pt after
            paragraph_format = current_paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
            paragraph_format.space_before = Pt(0)
            paragraph_format.left_indent = Inches(0.5)
            paragraph_format.first_line_indent = Inches(-0.5)
            
            # Add section header (bold)
            section_name = VALUES["{{abstract}}"][i]
            run = current_paragraph.add_run(section_name)
            apply_hardcoded_style(run, font_size=8, bold=True, spacing=0, scale=100)
            
            # Add a space between header and content
            current_paragraph.add_run(" ")
            
            # Add content (not bold) with percentage and symbol formatting
            content = VALUES["{{abstract}}"][i + 1]
            # Apply percentage formatting
            content = int_percent_to_float_percent(content)
            # Normalize spacing around = and ± symbols
            content = normalize_symbol_spacing(content)
            run = current_paragraph.add_run(content)
            apply_hardcoded_style(run, font_size=8, spacing=0, scale=100)
            
            # Move to next paragraph if not the last section
            if i + 2 < len(VALUES["{{abstract}}"]):
                # Create a new paragraph for the next section
                new_p_element = OxmlElement('w:p')
                current_p_element.addnext(new_p_element)
                current_paragraph = Paragraph(new_p_element, paragraph._parent)
                current_p_element = new_p_element


def _insert_keywords(paragraph: Paragraph) -> None:
    """Insert keywords with blue label."""
    run = paragraph.add_run(VALUES["{{keywords}}"][0])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=11, color="#2F5496", bold=True, spacing=0, scale=100)
    
    # Apply percentage formatting to keywords
    run = paragraph.add_run(VALUES["{{keywords}}"][1])
    apply_hardcoded_style(run, font_name="Times New Roman (Headings CS)", font_size=10, spacing=0, scale=100)


def _insert_content_section(paragraph: Paragraph, variable: str) -> None:
    """Insert content sections (intro, aim, methods, results, discussion, references)."""
    all_texts = VALUES[variable]
    
    # Check if the content is empty (just empty strings)
    has_content = any(text and text.strip() for text in all_texts)
    if not has_content:
        # If no content, just clear the placeholder and leave paragraph empty
        return
    
    # Use parent element to navigate and create paragraphs
    current_p_element = paragraph._p
    current_paragraph = paragraph
    
    for text_idx, text in enumerate(all_texts):
        # Check if this is a section header (ends with ":")
        if text and text.strip().endswith(":"):
            # Section header - add bold text
            run = current_paragraph.add_run(text)
            apply_hardcoded_style(run, font_size=10, bold=True, spacing=0, scale=100)
            
            # Move to next paragraph for content
            if text_idx < len(all_texts) - 1:
                # Always create a new paragraph to avoid overwriting section headers
                new_p_element = OxmlElement('w:p')
                current_p_element.addnext(new_p_element)
                current_paragraph = Paragraph(new_p_element, paragraph._parent)
                current_p_element = new_p_element
        else:
            # Content text - split by \n and create separate paragraphs
            if text and text.strip():
                # Split text by newlines (handle both \n\n and \n)
                text_paragraphs = []
                for chunk in text.split('\n\n'):
                    for para in chunk.split('\n'):
                        if para.strip():
                            text_paragraphs.append(para.strip())
                
                for para_idx, para_text in enumerate(text_paragraphs):
                    if para_text:  # Only process non-empty paragraphs
                        # Set justified alignment
                        current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        # Apply special formatting for intro, aim, methods (patients), and discussion
                        if variable in ["{{intro}}", "{{aim}}", "{{methods}}", "{{discussion}}"]:
                            process_formatted_text(current_paragraph, para_text)
                        else:
                            # Apply bracket, percentage, and symbol spacing formatting
                            formatted_text = para_text.replace('(', '[').replace(')', ']')
                            formatted_text = int_percent_to_float_percent(formatted_text)
                            formatted_text = normalize_symbol_spacing(formatted_text)
                            run = current_paragraph.add_run(formatted_text)
                            apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
                        
                        # Move to next paragraph if not the last text paragraph
                        if para_idx < len(text_paragraphs) - 1 or text_idx < len(all_texts) - 1:
                            # Always create a new paragraph instead of reusing existing ones
                            new_p_element = OxmlElement('w:p')
                            current_p_element.addnext(new_p_element)
                            current_paragraph = Paragraph(new_p_element, paragraph._parent)
                            if variable == "{{references}}":
                                current_paragraph.paragraph_format.left_indent = Inches(0.5)
                                current_paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                            else:
                                current_paragraph.paragraph_format.first_line_indent = Inches(0.5)
                            current_p_element = new_p_element


def _insert_tables(paragraph: Paragraph, doc: Optional[Document]) -> None:
    """Insert tables from VALUES."""
    table_num = 1
    current_caption = ""
    current_p_element = paragraph._p
    
    for _, item in enumerate(VALUES["{{tables}}"]):
        if item and str(item).strip():
            # Check if it's HTML table content
            if '<table' in str(item).lower():
                if doc:
                    # Create actual Word table from HTML
                    caption = current_caption if current_caption else f"Table {table_num}"
                    
                    # Create a temporary paragraph to pass positioning info
                    temp_para = doc.add_paragraph()
                    current_p_element.addnext(temp_para._p)
                    
                    # Process HTML table
                    last_para = process_table_content(temp_para, str(item), doc, table_num, caption)

                    # Update position tracking: continue insertion after the whole table block
                    current_p_element = (last_para._p if last_para is not None else temp_para._p)
                    
                    # Increment table counter and reset caption
                    table_num += 1
                    current_caption = ""
            
            # Check if it's a caption (starts with "Table X:")
            elif str(item).strip().startswith("Table ") and ':' in str(item):
                # Save the caption for the next table
                current_caption = str(item).strip()
            
            # Skip "\n\n" separators
            elif str(item).strip() == "\n\n":
                continue


def _insert_figures(paragraph: Paragraph, doc: Optional[Document]) -> None:
    """Insert figures from VALUES."""
    figure_num = 1
    current_caption = ""
    current_p_element = paragraph._p
    
    for _, item in enumerate(VALUES["{{figures}}"]):
        if item and str(item).strip():
            # Check if it's image content
            if '<img' in str(item).lower():
                if doc:
                    # Create a new paragraph for the image
                    img_para = doc.add_paragraph()
                    current_p_element.addnext(img_para._p)
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Process the image
                    process_image_content(img_para, str(item), figure_num, current_caption)
                    
                    # Add caption below image if we have one
                    if current_caption:
                        caption_para = doc.add_paragraph()
                        img_para._p.addnext(caption_para._p)
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Set caption spacing: 6pt before and 6pt after
                        caption_para.paragraph_format.space_after = Pt(6)
                        caption_para.paragraph_format.space_before = Pt(6)
                        
                        # Parse caption: "Figure X:" in bold, rest normal (same as tables)
                        if ':' in current_caption:
                            colon_idx = current_caption.index(':')
                            label_part = current_caption[:colon_idx + 1]
                            desc_part = current_caption[colon_idx + 1:].strip()
                            
                            label_run = caption_para.add_run(label_part)
                            apply_hardcoded_style(label_run, font_size=10, bold=True, spacing=0, scale=100)
                            
                            if desc_part:
                                caption_para.add_run(" ")
                                desc_run = caption_para.add_run(desc_part)
                                apply_hardcoded_style(desc_run, font_size=10, spacing=0, scale=100)
                        else:
                            caption_run = caption_para.add_run(current_caption)
                            apply_hardcoded_style(caption_run, font_size=10, bold=True, spacing=0, scale=100)
                        
                        current_p_element = caption_para._p
                    else:
                        current_p_element = img_para._p
                    
                    # Increment counter and reset caption
                    figure_num += 1
                    current_caption = ""
            
            # Check if it's a caption (starts with "Figure X:")
            elif str(item).strip().startswith("Figure ") and ':' in str(item):
                # Save the caption for the next figure
                current_caption = str(item).strip()
            
            # Skip "\n\n" separators
            elif str(item).strip() == "\n\n":
                continue
