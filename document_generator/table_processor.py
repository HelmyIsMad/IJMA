"""
Table processing utilities for document generation.
Handles HTML table parsing, Word table creation, and cell merging.
"""
from typing import List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
import re
from .style_handler import apply_hardcoded_style


def _normalize_symbol_spacing_preserve_newlines(text: str) -> str:
    """Normalize spacing around '=' and '±' without destroying newlines."""
    # Apply per-line so we don't collapse \n into spaces.
    lines = text.split('\n')
    normalized_lines: List[str] = []
    for line in lines:
        # Remove all spaces around = and ±
        line = re.sub(r'\s*=\s*', '=', line)
        line = re.sub(r'\s*±\s*', '±', line)

        # Add exactly one space around
        line = re.sub(r'=', ' = ', line)
        line = re.sub(r'±', ' ± ', line)

        # Clean up only spaces/tabs (not newlines)
        line = re.sub(r'[ \t]+', ' ', line)
        normalized_lines.append(line.strip())

    return '\n'.join(normalized_lines)


def process_table_content(paragraph: Paragraph, html_content: str, doc: Optional[Document] = None, 
                          table_num: int = 1, caption: str = "") -> None:
    """
    Process HTML table content and create actual Word table.
    
    Args:
        paragraph: The paragraph to insert the table after
        html_content: HTML content containing the table
        doc: The document object
        table_num: Table number for caption
        caption: Table caption text
    """
    try:
        from bs4 import BeautifulSoup
        
        # Parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        table_tag = soup.find('table')
        
        if table_tag and doc:
            # Extract table data
            rows = table_tag.find_all('tr')
            if not rows:
                # Fallback to text
                _fallback_to_text(paragraph, html_content)
                return

            # Build 2D array from HTML table, handling rowspan and colspan
            table_data, _ = _parse_html_table_to_grid(rows)
            
            # Create actual Word table with merges
            if table_data:
                # Create Word table with HTML structure and merges
                _create_word_table_from_html(paragraph, table_data, doc, rows, table_num, caption)
                return
                
    except ImportError:
        # BeautifulSoup not available, fallback to regex
        pass
    except Exception as e:
        # Error processing table, fall through to text fallback
        pass
    
    # Fallback: strip HTML and add as text
    _fallback_to_text(paragraph, html_content)


def _parse_html_table_to_grid(rows) -> tuple:
    """
    Parse HTML table rows into a 2D grid, handling rowspan and colspan.
    
    Args:
        rows: List of HTML table row elements
        
    Returns:
        Tuple of (table_data, max_cols)
    """
    # Determine the maximum number of columns
    max_cols = 0
    for row in rows:
        cells = row.find_all(['td', 'th'])
        col_count = sum(int(cell.get('colspan', 1)) for cell in cells)
        max_cols = max(max_cols, col_count)
    
    # Build a grid to handle rowspan and colspan
    grid = []
    for row_idx, row in enumerate(rows):
        if row_idx >= len(grid):
            grid.append([None] * max_cols)
        
        cells = row.find_all(['td', 'th'])
        col_idx = 0
        
        for cell in cells:
            # Find the next available column
            while col_idx < max_cols and grid[row_idx][col_idx] is not None:
                col_idx += 1
            
            if col_idx >= max_cols:
                break
            
            # Get cell attributes
            cell_text = cell.get_text(separator='\n').strip()
            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))
            
            # Fill the grid for this cell and its spans
            for r in range(rowspan):
                # Ensure we have enough rows
                while len(grid) <= row_idx + r:
                    grid.append([None] * max_cols)
                
                for c in range(colspan):
                    if col_idx + c < max_cols:
                        # Only put text in the first cell, empty string for merged cells
                        if r == 0 and c == 0:
                            grid[row_idx + r][col_idx + c] = cell_text
                        else:
                            grid[row_idx + r][col_idx + c] = ""
            
            col_idx += colspan
    
    # Convert grid to table_data (replace None with empty string)
    table_data = [[cell if cell is not None else "" for cell in row] for row in grid]
    return table_data, max_cols




def _apply_table_style(word_table) -> None:
    """Apply a table style to the Word table."""
    style_applied = False
    style_names = ['Grid Table 4 Accent 1', 'GridTable4-Accent1', 'Grid Table 4 - Accent 1', 'Table Grid']
    
    for style_name in style_names:
        try:
            word_table.style = style_name
            style_applied = True
            break
        except:
            continue
    
    if not style_applied:
        word_table.style = 'Table Grid'


def _apply_table_formatting_options(word_table) -> None:
    """Enable table formatting options using direct XML manipulation."""
    try:
        tbl = word_table._element
        tblPr = tbl.tblPr
        
        # Create or update tblLook element
        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is None:
            tblLook = OxmlElement('w:tblLook')
            tblPr.append(tblLook)
        
        # Set table look attributes for banded rows/columns and header
        tblLook.set(qn('w:firstRow'), '1')      # Header row
        tblLook.set(qn('w:lastRow'), '0')       # No last row
        tblLook.set(qn('w:firstColumn'), '0')   # No first column  
        tblLook.set(qn('w:lastColumn'), '0')    # No last column
        tblLook.set(qn('w:noHBand'), '0')       # Enable horizontal banding
        tblLook.set(qn('w:noVBand'), '0')       # Enable vertical banding
        
    except Exception as e:
        # Could not apply table formatting, continue
        pass


def _add_table_caption(doc: Document, paragraph: Paragraph, word_table, table_num: int, caption: str) -> Paragraph:
    """Add a caption above the table."""
    caption_para = doc.add_paragraph()
    paragraph._p.addnext(caption_para._p)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set caption spacing: 0pt after (no space between caption and table)
    caption_para.paragraph_format.space_after = Pt(0)
    caption_para.paragraph_format.space_before = Pt(6)
    
    # Use provided caption or default
    table_caption = caption if caption else f"Table {table_num}"
    
    # Parse caption to separate "Table X:" from description
    if ':' in table_caption:
        colon_idx = table_caption.index(':')
        label_part = table_caption[:colon_idx + 1]
        desc_part = table_caption[colon_idx + 1:].strip()
        
        # Add label in bold
        label_run = caption_para.add_run(label_part)
        apply_hardcoded_style(label_run, font_size=10, bold=True, spacing=0, scale=100)
        
        # Add space and description in regular font
        if desc_part:
            caption_para.add_run(" ")
            desc_run = caption_para.add_run(desc_part)
            apply_hardcoded_style(desc_run, font_size=10, spacing=0, scale=100)
    else:
        # No colon, make entire caption bold
        caption_run = caption_para.add_run(table_caption)
        apply_hardcoded_style(caption_run, font_size=10, bold=True, spacing=0, scale=100)
    
    # Move table after caption
    caption_para._p.addnext(word_table._element)
    
    return caption_para


def _add_table_spacing(doc: Document, word_table) -> None:
    """Add spacing paragraph after the table."""
    spacing_para = doc.add_paragraph()
    word_table._element.addnext(spacing_para._p)
    spacing_para.paragraph_format.space_after = Pt(6)
    spacing_para.paragraph_format.space_before = Pt(0)


def _create_word_table_from_html(paragraph: Paragraph, table_data: list, doc: Document,
                                 html_rows, table_num: int, caption: str) -> None:
    """
    Create Word table from HTML table data with proper formatting and merges.
    
    Args:
        paragraph: The paragraph to insert the table after
        table_data: 2D array of table data
        doc: The document object
        html_rows: Original HTML rows for merge information
        table_num: Table number for caption
        caption: Table caption text
    """
    try:
        if not table_data or len(table_data) == 0:
            return
            
        # Determine table size
        max_cols = max(len(row) for row in table_data) if table_data else 0
        
        if max_cols > 0:
            # Create Word table
            word_table = doc.add_table(rows=len(table_data), cols=max_cols)
            paragraph._p.addnext(word_table._element)
            
            # Center the table
            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set table style
            _apply_table_style(word_table)
            
            # Add caption
            caption_para = _add_table_caption(doc, paragraph, word_table, table_num, caption)
            
            # Add spacing after table
            _add_table_spacing(doc, word_table)
            
            # Parse HTML to get merge information
            merge_map = _parse_merge_information(html_rows)
            
            # First pass: Apply all cell merges
            merged_cells = _apply_cell_merges(word_table, table_data, merge_map)
            
            # Second pass: Fill table with data and apply styling
            _fill_and_style_table(word_table, table_data, merged_cells)
            
    except Exception as e:
        # If anything fails, fall back to text
        _fallback_to_text(paragraph, str(table_data))


def _parse_merge_information(html_rows) -> dict:
    """Parse HTML rows to extract cell merge information."""
    merge_map = {}
    try:
        row_idx = 0
        for html_row in html_rows:
            cells = html_row.find_all(['td', 'th'])
            col_idx = 0
            for cell in cells:
                # Find the next available column
                while (row_idx, col_idx) in merge_map:
                    col_idx += 1
                
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))
                
                if rowspan > 1 or colspan > 1:
                    merge_map[(row_idx, col_idx)] = {
                        'rowspan': rowspan,
                        'colspan': colspan
                    }
                
                # Mark spanned cells as occupied
                for r in range(rowspan):
                    for c in range(colspan):
                        if r > 0 or c > 0:
                            merge_map[(row_idx + r, col_idx + c)] = 'spanned'
                
                col_idx += colspan
            row_idx += 1
    except Exception as e:
        # Could not parse merge information, continue without merges
        pass
    
    return merge_map


def _apply_cell_merges(word_table, table_data: list, merge_map: dict) -> set:
    """Apply cell merges to the Word table and return set of merged cells."""
    merged_cells = set()
    
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_value in enumerate(row_data):
            # Skip cells that are marked as spanned
            if (row_idx, col_idx) in merge_map and merge_map[(row_idx, col_idx)] == 'spanned':
                merged_cells.add((row_idx, col_idx))
                continue
            
            # Apply cell merges if needed
            if (row_idx, col_idx) in merge_map and merge_map[(row_idx, col_idx)] != 'spanned':
                merge_info = merge_map[(row_idx, col_idx)]
                rowspan = merge_info['rowspan']
                colspan = merge_info['colspan']
                
                # Merge cells
                if rowspan > 1 or colspan > 1:
                    try:
                        cell = word_table.rows[row_idx].cells[col_idx]
                        end_row = min(row_idx + rowspan - 1, len(word_table.rows) - 1)
                        end_col = min(col_idx + colspan - 1, len(word_table.rows[row_idx].cells) - 1)
                        cell.merge(word_table.rows[end_row].cells[end_col])
                        
                        # Mark merged cells
                        for r in range(rowspan):
                            for c in range(colspan):
                                if r > 0 or c > 0:
                                    merged_cells.add((row_idx + r, col_idx + c))
                    except Exception as e:
                        # Could not merge cell, continue
                        pass
    
    return merged_cells


def _fill_and_style_table(word_table, table_data: list, merged_cells: set) -> None:
    """Fill table with data and apply styling."""
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_value in enumerate(row_data):
            # Skip cells that have been merged into another cell
            if (row_idx, col_idx) in merged_cells:
                continue
            
            if col_idx < len(word_table.rows[row_idx].cells):
                cell = word_table.rows[row_idx].cells[col_idx]
                
                # Apply formatting
                cell_text = str(cell_value) if cell_value else ""
                cell_text = cell_text.replace('(', '[').replace(')', ']')
                # Apply percentage formatting and normalize symbol spacing
                from .text_formatter import apply_percentage_formatting
                cell_text = apply_percentage_formatting(cell_text)
                cell_text = _normalize_symbol_spacing_preserve_newlines(cell_text)
                cell.text = cell_text
                
                # Style the cell
                for paragraph_in_cell in cell.paragraphs:
                    paragraph_in_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph_in_cell.runs:
                        apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
                
                # Make header row bold
                if row_idx == 0:
                    for paragraph_in_cell in cell.paragraphs:
                        for run in paragraph_in_cell.runs:
                            apply_hardcoded_style(run, font_size=10, bold=True, spacing=0, scale=100)
                
                # Set vertical alignment
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = OxmlElement('w:vAlign')
                    tcPr.append(vAlign)
                vAlign.set(qn('w:val'), 'center')


def _fallback_to_text(paragraph: Paragraph, html_content: str) -> None:
    """Fallback: strip HTML and add as text."""
    # Preserve explicit line breaks before stripping tags
    clean_text = re.sub(r'(?i)<\s*br\s*/?>', '\n', html_content)
    clean_text = re.sub(r'</\s*(div|p|tr)\s*>', '\n', clean_text, flags=re.I)
    clean_text = re.sub(r'<[^>]+>', ' ', clean_text)

    # Collapse spaces/tabs but keep newlines
    clean_text = re.sub(r'[ \t\r\f\v]+', ' ', clean_text)
    clean_text = re.sub(r'\n\s*\n+', '\n', clean_text)
    clean_text = clean_text.strip()
    if clean_text:
        run = paragraph.add_run(clean_text)
        apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
        paragraph.add_run("\n")
