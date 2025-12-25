"""
Main document processor - coordinates document generation workflow.
Simplified version that delegates to specialized modules.
"""
from typing import List
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.table import Table
from .value_inserter import insert_values
from .text_formatter import apply_global_formatting_rules

def iter_paragraphs(parent):
    for child in parent.element.body.iter():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def get_all_paragraphs(doc):
    paragraphs = []

    # body
    for p in doc.paragraphs:
        paragraphs.append(p)

    # tables (body)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    # headers & footers
    for section in doc.sections:
        for hf in (section.header, section.footer):
            paragraphs.extend(hf.paragraphs)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)

    return paragraphs

def process_document(doc: Document, variables: List[str]) -> None:
    """
    Process entire document by replacing placeholders with formatted content.
    
    Args:
        doc: The Word document to process
        variables: List of placeholder variables to replace
    """
    # Process all paragraphs in the document
    paragraphs = get_all_paragraphs(doc)
    process_paragraphs(paragraphs, variables, doc)
    
    # Apply global formatting rules to all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text and not any(var in paragraph.text for var in variables):
            apply_global_formatting_rules(paragraph)

def process_paragraphs(paragraphs: List[Paragraph], variables: List[str], doc: Document = None) -> None:
    """
    Process a list of paragraphs, replacing template variables with values.
    
    Args:
        paragraphs: List of paragraphs to process
        variables: List of placeholder variables to search for
        doc: The document object (needed for tables/figures)
    """
    for paragraph in paragraphs:
        for variable in variables:
            if variable in paragraph.text:
                insert_values(paragraph, variable, doc)
                break