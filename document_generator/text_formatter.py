"""
Text formatting utilities for document processing.
Handles text transformations, special character replacements, and formatting rules.
"""
from typing import List
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import re
from .style_handler import apply_hardcoded_style


def normalize_symbol_spacing(text: str) -> str:
    """
    Ensure "=" and "±" have exactly one space before and after them.
    
    Args:
        text: The text to normalize
        
    Returns:
        Text with normalized spacing around = and ± symbols
    """
    # Remove all spaces around = and ±
    text = re.sub(r'\s*=\s*', '=', text)
    text = re.sub(r'\s*±\s*', '±', text)
    
    # Add exactly one space before and after
    text = re.sub(r'=', ' = ', text)
    text = re.sub(r'±', ' ± ', text)
    
    # Clean up any double spaces that might have been created
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


def apply_text_formatting_rules(paragraph: Paragraph) -> None:
    """
    Apply special formatting rules to paragraph text.
    
    Args:
        paragraph: The paragraph to format
    """
    full_text = paragraph.text
    paragraph.clear()
    process_formatted_text(paragraph, full_text)


def process_formatted_text(paragraph: Paragraph, text: str) -> None:
    """
    Process text and apply formatting rules:
    - Replace round brackets with square brackets
    - Convert integer percentages to float format (10% -> 10.0%)
    - Normalize spacing around = and ± symbols
    - Apply superscript formatting to bracketed numbers
    - Apply italic formatting to "et al"
    
    Args:
        paragraph: The paragraph to add formatted text to
        text: The text to process and format
    """
    # Replace round brackets with square brackets
    text = text.replace('(', '[').replace(')', ']')
    
    # Convert integer percentages to float format (10% -> 10.0%)
    text = apply_percentage_formatting(text)
    
    # Normalize spacing around = and ± symbols
    text = normalize_symbol_spacing(text)
    
    # Split text into segments for different formatting
    segments = []
    current_pos = 0
    
    # Find numbers in square brackets (originally parentheses) for superscript
    bracket_pattern = r'\[(\d+)\]'
    for match in re.finditer(bracket_pattern, text):
        # Add text before match
        if match.start() > current_pos:
            segments.append({
                'text': text[current_pos:match.start()],
                'format': 'normal'
            })
        
        # Add the bracketed number with superscript formatting
        segments.append({
            'text': match.group(0),  # Full match including brackets
            'format': 'superscript_bold'
        })
        
        current_pos = match.end()
    
    # Add remaining text
    if current_pos < len(text):
        segments.append({
            'text': text[current_pos:],
            'format': 'normal'
        })
    
    # Process segments and apply formatting
    for segment in segments:
        segment_text = segment['text']
        
        # Further process for "et al" formatting
        if segment['format'] == 'normal':
            # Split on "et al" for special formatting with author name
            # Pattern: captures word before "et al" (author name), "et al", and rest
            parts = re.split(r'(\S+)(\s+)(et al\b)', segment_text)
            
            i = 0
            while i < len(parts):
                part = parts[i]
                
                # Check if this is followed by whitespace and "et al"
                if i + 3 < len(parts) and parts[i + 2] == 'et al':
                    # This is the author name before "et al" - make it bold
                    author_name = parts[i]
                    whitespace = parts[i + 1]
                    et_al = parts[i + 2]
                    
                    if author_name:
                        run = paragraph.add_run(author_name)
                        apply_hardcoded_style(run, font_size=10, bold=True, spacing=0, scale=100)
                    
                    if whitespace:
                        paragraph.add_run(whitespace)
                    
                    # "et al" - bold and italic
                    run = paragraph.add_run(et_al)
                    apply_hardcoded_style(run, font_size=10, bold=True, italic=True, spacing=0, scale=100)
                    
                    i += 3  # Skip the parts we just processed
                elif part == 'et al':
                    # Standalone "et al" without captured author name
                    run = paragraph.add_run(part)
                    apply_hardcoded_style(run, font_size=10, bold=True, italic=True, spacing=0, scale=100)
                    i += 1
                elif part:
                    # Regular text
                    run = paragraph.add_run(part)
                    apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
                    i += 1
                else:
                    i += 1
        else:
            # Apply special formatting
            run = paragraph.add_run(segment_text)
            if segment['format'] == 'superscript_bold':
                apply_hardcoded_style(run, font_size=10, bold=True, superscript=True, spacing=0, scale=100)
            else:
                apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)


def apply_percentage_formatting(text: str) -> str:
    """
    Convert integer percentages to float format (10% -> 10.0%).
    
    Args:
        text: The text to format
        
    Returns:
        Formatted text with decimal percentages
    """
    return re.sub(
        r'(?<!\.)\d+(?=%)',
        lambda m: f"{m.group(0)}.0",
        text
    )


def apply_bracket_conversion(text: str) -> str:
    """
    Replace round brackets with square brackets.
    
    Args:
        text: The text to format
        
    Returns:
        Text with square brackets
    """
    return text.replace('(', '[').replace(')', ']')


def apply_global_formatting_rules(paragraph: Paragraph) -> None:
    """
    Apply global formatting rules to a paragraph (brackets and percentages).
    
    Args:
        paragraph: The paragraph to format
    """
    full_text = paragraph.text
    
    # Apply transformations
    formatted_text = apply_bracket_conversion(full_text)
    formatted_text = apply_percentage_formatting(formatted_text)
    
    # Only update if text changed
    if formatted_text != full_text:
        paragraph.clear()
        run = paragraph.add_run(formatted_text)
        apply_hardcoded_style(run, font_size=10, spacing=0, scale=100)
